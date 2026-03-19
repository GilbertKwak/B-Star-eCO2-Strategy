#!/usr/bin/env python3
"""
B★ eCO₂ Strategy — DOCX 병합 스크립트 v2.2
개선: 표 진짜 DOCX 변환, 블릿 포인트 개선, 심평 머리말 스타일
사용법:
  python merge_to_docx.py --vol 1 --lang KR
  python merge_to_docx.py --vol all --lang all
"""

import argparse
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx 미설치: pip install python-docx")
    exit(1)

VOLUME_MAP = {
    "1": {
        "KR": {"dir": "Vol1_Technical/KR",
               "files": ["Ch01_KR.md","Ch02_KR.md","Ch03_KR.md","Ch04_KR.md"],
               "title": "B★ eCO₂ 전략 보고서 — Vol.1 기술편",
               "subtitle": "응용처 분석 · 열역학 · 열원 설계 · 압축기PCHE"},
        "EN": {"dir": "Vol1_Technical/EN",
               "files": ["Ch01_EN.md","Ch02_EN.md","Ch03_EN.md","Ch04_EN.md"],
               "title": "B★ eCO₂ Strategy Report — Vol.1 Technical",
               "subtitle": "Application Analysis · Thermodynamics · Heat Source · Compressor PCHE"}
    },
    "2": {
        "KR": {"dir": "Vol2_Business/KR",
               "files": ["Ch05_KR.md","Ch06_KR.md","Ch07_KR.md","Ch08_KR.md","Ch09_KR.md"],
               "title": "B★ eCO₂ 전략 보고서 — Vol.2 사업편",
               "subtitle": "제어전략 · 안전규제 · 사업화 · 재무 · IP특허"},
        "EN": {"dir": "Vol2_Business/EN",
               "files": ["Ch05_EN.md","Ch06_EN.md","Ch07_EN.md","Ch08_EN.md","Ch09_EN.md"],
               "title": "B★ eCO₂ Strategy Report — Vol.2 Business",
               "subtitle": "Control · Safety · Commercialization · Finance · IP"}
    },
    "3": {
        "KR": {"dir": "Vol3_Appendix/KR",
               "files": ["AppA_KR.md","AppB_KR.md","AppC_KR.md"],
               "title": "B★ eCO₂ 전략 보고서 — Vol.3 부록편",
               "subtitle": "차트 · 데이터 테이블 · 용어집 · 참고문헌"},
        "EN": {"dir": "Vol3_Appendix/EN",
               "files": ["AppA_EN.md","AppB_EN.md","AppC_EN.md"],
               "title": "B★ eCO₂ Strategy Report — Vol.3 Appendix",
               "subtitle": "Charts · Data Tables · Glossary · References"}
    }
}


def set_cell_bg(cell, hex_color):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_cover_page(doc, title, subtitle, lang="KR"):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(subtitle).font.size = Pt(14)
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vt = ("버전: v2.0 | 날짜: 2026-03-19 | 분류: 내부용"
          if lang == "KR" else
          "Version: v2.0 | Date: 2026-03-19 | Classification: Confidential")
    run3 = p3.add_run(vt)
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_page_break()


def add_md_table(doc, table_lines):
    """Parse markdown table lines into a real DOCX table"""
    rows = []
    for line in table_lines:
        if re.match(r'^\|[-:\s|]+\|\s*$', line):
            continue  # skip separator row
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    # Normalize row lengths
    rows = [r + [''] * (max_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            # Style header row
            if i == 0:
                set_cell_bg(cell, 'E8F0FE')
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(10)
            else:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


def parse_md_and_add(doc, md_path):
    if not os.path.exists(md_path):
        print(f"  [SKIP] 미존재: {md_path}")
        p = doc.add_paragraph(f"[⚠️ {os.path.basename(md_path)} 미배치 — 소스 파일 Push 후 자동 재빌드]")
        p.runs[0].italic = True
        doc.add_page_break()
        return

    print(f"  → 병합: {md_path}")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    table_buffer = []
    in_table = False
    in_code = False

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        # Code block toggle
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            doc.add_paragraph(line, style='No Spacing')
            continue

        # Table detection
        is_table_line = line.strip().startswith('|')
        if is_table_line:
            table_buffer.append(line)
            in_table = True
            continue
        elif in_table:
            add_md_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(style='Quote')
            p.add_run(line[2:]).italic = True
        # Horizontal rule
        elif line.strip().startswith("---"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.add_run(re.sub(r'^\d+\.\s', '', line))
        # Normal paragraph
        elif line.strip():
            # Handle **bold** inline
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        # Blank line
        elif not line.strip():
            doc.add_paragraph()

    # Flush remaining table buffer
    if table_buffer:
        add_md_table(doc, table_buffer)

    doc.add_page_break()


def build_volume(vol_num, lang, output_dir="output"):
    config = VOLUME_MAP[str(vol_num)][lang]
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.0)

    add_cover_page(doc, config["title"], config["subtitle"], lang)

    toc_label = "목차" if lang == "KR" else "Table of Contents"
    toc_hint  = ("[Word에서 '참조 > 목차 업데이트' 실행]"
                 if lang == "KR" else
                 "[Update in Word: References > Update Table]")
    doc.add_heading(toc_label, level=1)
    doc.add_paragraph(toc_hint)
    doc.add_page_break()

    base_dir = Path(__file__).parent.parent
    for md_file in config["files"]:
        md_path = base_dir / config["dir"] / md_file
        parse_md_and_add(doc, str(md_path))

    vol_label = {"1": "기술편" if lang=="KR" else "Technical",
                 "2": "사업편" if lang=="KR" else "Business",
                 "3": "부록편" if lang=="KR" else "Appendix"}
    filename = f"BStar_eCO2_Vol{vol_num}_{vol_label[str(vol_num)]}_v2.0_{lang}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) / 1024
    print(f"  ✅ 저장: {out_path} ({size_kb:.0f} KB)")
    return out_path


def main():
    parser = argparse.ArgumentParser(description="B★ eCO₂ DOCX 병합 빌더 v2.2")
    parser.add_argument("--vol",  default="all", help="볼륨: 1, 2, 3, all")
    parser.add_argument("--lang", default="all", help="언어: KR, EN, all")
    parser.add_argument("--out",  default="output", help="출력 디렉토리")
    args = parser.parse_args()

    vols  = ["1","2","3"] if args.vol  == "all" else [args.vol]
    langs = ["KR","EN"]   if args.lang == "all" else [args.lang.upper()]

    print(f"\n🚀 B★ eCO₂ DOCX 병합 v2.2 — Vol:{vols} / Lang:{langs} / Out:{args.out}\n")

    results = []
    for vol in vols:
        for lang in langs:
            print(f"\n[Vol.{vol}-{lang}] 빌드 시작...")
            try:
                path = build_volume(vol, lang, args.out)
                results.append(path)
            except Exception as e:
                print(f"  [오류] Vol.{vol}-{lang}: {e}")

    print(f"\n✅ 완료 — {len(results)}개 DOCX 생성")
    for r in results:
        kb = os.path.getsize(r)/1024 if os.path.exists(r) else 0
        print(f"   {r} ({kb:.0f} KB)")


if __name__ == "__main__":
    main()
